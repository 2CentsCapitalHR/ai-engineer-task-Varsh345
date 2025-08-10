import docx

def annotate_docx(input_file, issues, output_file):
    doc = docx.Document(input_file)
    doc.add_paragraph("\n--- Review Comments ---")
    for issue in issues:
        doc.add_paragraph(f"{issue['document']} - {issue['issue']} | Suggestion: {issue['suggestion']}")
    doc.save(output_file)
